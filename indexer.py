import os
import sys
from dotenv import load_dotenv

# Carga de variables de entorno
load_dotenv()

from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# ARCHIVOS FUENTE EN LA CARPETA ACTUAL
PDF_PATH = "Base_de_Conocimiento_ISTCGE_Instituo.pdf"
DOCX_PATH = "Base_Entrenamiento_NIA_ISTCGE_ASCEND.docx"
DB_FAISS_PATH = "vectorstore/db_faiss"

def load_documents():
    all_documents = []
    
    # 1. Cargar PDF si existe
    if os.path.exists(PDF_PATH):
        print(f"📄 Cargando PDF: {PDF_PATH}...")
        try:
            pdf_loader = PyPDFLoader(PDF_PATH)
            pdf_docs = pdf_loader.load()
            print(f"   -> {len(pdf_docs)} páginas cargadas del PDF.")
            all_documents.extend(pdf_docs)
        except Exception as e:
            print(f"⚠️ Error cargando PDF con PyPDFLoader: {e}")
            try:
                import pypdf
                from langchain_core.documents import Document
                reader = pypdf.PdfReader(PDF_PATH)
                pdf_docs = [
                    Document(page_content=page.extract_text() or "", metadata={"source": PDF_PATH, "page": i})
                    for i, page in enumerate(reader.pages)
                ]
                print(f"   -> {len(pdf_docs)} páginas extraídas mediante pypdf.")
                all_documents.extend(pdf_docs)
            except Exception as ex:
                print(f"❌ Falló extracción de PDF: {ex}")
    else:
        print(f"⚠️ Advertencia: No se encontró el archivo PDF en: {PDF_PATH}")

    # 2. Cargar DOCX si existe
    if os.path.exists(DOCX_PATH):
        print(f"📝 Cargando DOCX: {DOCX_PATH}...")
        try:
            docx_loader = Docx2txtLoader(DOCX_PATH)
            docx_docs = docx_loader.load()
            print(f"   -> {len(docx_docs)} secciones cargadas del DOCX.")
            all_documents.extend(docx_docs)
        except Exception as e:
            print(f"⚠️ Error cargando DOCX con Docx2txtLoader: {e}")
            try:
                import docx
                from langchain_core.documents import Document
                doc = docx.Document(DOCX_PATH)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                docx_docs = [Document(page_content=full_text, metadata={"source": DOCX_PATH})]
                print(f"   -> DOCX extraído exitosamente con python-docx ({len(full_text)} caracteres).")
                all_documents.extend(docx_docs)
            except Exception as ex:
                print(f"❌ Falló extracción de DOCX: {ex}")
    else:
        print(f"⚠️ Advertencia: No se encontró el archivo DOCX en: {DOCX_PATH}")

    return all_documents

def create_vector_db(use_local_embeddings=False):
    print("🚀 Iniciando indexado y vectorización de NIA - Asistente ISTCGE...")
    
    docs = load_documents()
    if not docs:
        print("❌ ERROR: No se encontraron documentos para procesar.")
        return False
        
    total_chars = sum(len(d.page_content) for d in docs)
    print(f"📊 Total de contenido extraído: {total_chars} caracteres en {len(docs)} elementos base.")
    
    # División en fragmentos (Chunks)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=120,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    chunks = text_splitter.split_documents(docs)
    chunks = [c for c in chunks if c.page_content.strip()]
    print(f"✂️ Texto dividido en {len(chunks)} fragmentos optimizados.")
    
    # Configurar modelo de embeddings
    api_key = os.getenv("GEMINI_API_KEY")
    
    if not use_local_embeddings and api_key:
        print("🔑 Usando embeddings de Google AI Studio (models/gemini-embedding-001)...")
        try:
            embeddings = GoogleGenerativeAIEmbeddings(
                model="models/gemini-embedding-001",
                google_api_key=api_key
            )
            embeddings.embed_query("Prueba de conexión")
        except Exception as e:
            print(f"⚠️ Error con Google Embeddings: {e}. Cambiando a modelo local multilingual...")
            use_local_embeddings = True
    else:
        use_local_embeddings = True
        
    if use_local_embeddings:
        print("🤖 Cargando modelo de embeddings local: sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2...")
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
        )
        
    # Crear y guardar índice FAISS
    print("📦 Generando base vectorial en FAISS...")
    vector_store = FAISS.from_documents(chunks, embeddings)
    
    os.makedirs(os.path.dirname(DB_FAISS_PATH), exist_ok=True)
    vector_store.save_local(DB_FAISS_PATH)
    
    print(f"✅ ¡ÉXITO TOTAL! Base de datos vectorial guardada en: {DB_FAISS_PATH}")
    print(f"   Archivos creados: {os.path.join(DB_FAISS_PATH, 'index.faiss')} e {os.path.join(DB_FAISS_PATH, 'index.pkl')}")
    return True

if __name__ == "__main__":
    use_local = "--local" in sys.argv
    create_vector_db(use_local_embeddings=use_local)
